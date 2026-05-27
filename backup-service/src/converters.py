"""Conversión de emails a MD, PDF y DOCX."""

import html as html_lib
import io
import logging
import re
from datetime import datetime, timezone
from typing import Optional

logger = logging.getLogger(__name__)
UTC = timezone.utc


def _slugify(text: str, max_len: int = 40) -> str:
    if not text:
        return "email"
    safe = "".join(c if c.isalnum() or c in " -_." else "_" for c in text)
    safe = re.sub(r"\s+", " ", safe).strip()
    return safe[:max_len].strip()


def _strip_html(text: str) -> str:
    clean = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    clean = re.sub(r"<[^>]+>", "", clean)
    clean = html_lib.unescape(clean)
    return clean


def _plain_body(content: dict) -> str:
    if content.get("text_plain"):
        return content["text_plain"]
    if content.get("text_html"):
        return _strip_html(content["text_html"])
    return "(sin contenido de texto)"


def _build_filename(subject: str, email_date: str, ext: str) -> str:
    slug = _slugify(subject)
    now = datetime.now(UTC).strftime("%Y%m%d")
    edate = ""
    if email_date:
        try:
            from email.utils import parsedate_to_datetime
            edate = parsedate_to_datetime(email_date).strftime("%Y%m%d")
        except Exception:
            pass
    if edate:
        return f"{slug} -restored{now} -mailde{edate}.{ext}"
    return f"{slug} -restored{now}.{ext}"


# =============================================================================
# Markdown
# =============================================================================

def email_to_markdown_bytes(content: dict, subject: str, email_date: str) -> bytes:
    h = content.get("headers", {})
    lines = [
        "---",
        f"From: {h.get('from', '')}",
        f"To: {h.get('to', '')}",
        f"Date: {h.get('date', '')}",
        f"Subject: {h.get('subject', subject)}",
        f"Message-ID: {h.get('message_id', '')}",
        "---",
        "",
        _plain_body(content),
    ]
    return ("\n".join(lines) + "\n").encode("utf-8")

def email_to_markdown(path: str) -> Optional[bytes]:
    from .maildir import get_email_content
    content = get_email_content(path)
    if not content:
        return None
    h = content.get("headers", {})
    return email_to_markdown_bytes(content, h.get("subject", ""), h.get("date", ""))


# =============================================================================
# PDF (reportlab)
# =============================================================================

def _md_text_to_story(text: str, styles) -> list:
    """Convierte texto (puede ser markdown o plain) en elementos reportlab."""
    from reportlab.platypus import Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm

    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=16, leading=20, spaceAfter=10, textColor="#222222")
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=13, leading=17, spaceAfter=8, textColor="#0d6efd")
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=11, leading=15, spaceAfter=6, textColor="#333333")
    body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=4)
    code = ParagraphStyle("Code", parent=styles["Normal"], fontSize=8.5, leading=12, fontName="Courier",
                          backColor="#f5f5f5", spaceAfter=4, leftIndent=8, rightIndent=8)
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=20, bulletIndent=10, spaceBefore=1, spaceAfter=1)

    story = []
    in_code = False
    in_table = False
    table_rows = []

    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                in_code = False
                continue
            else:
                in_code = True
                continue
        if in_table and "---" in stripped and "|" in stripped:
            continue
        if stripped.startswith("| ") and not in_table:
            in_table = True
            continue
        if in_table and not stripped.startswith("|"):
            if table_rows:
                story.append(Spacer(1, 4))
                for row in table_rows:
                    cells = [c.strip() for c in row.split("|")[1:-1]]
                    story.append(Paragraph("  " + "  —  ".join(cells), body))
                story.append(Spacer(1, 4))
            table_rows = []
            in_table = False
        if in_table:
            table_rows.append(stripped)
            continue
        if in_code:
            story.append(Paragraph(stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), code))
            continue

        if not stripped:
            story.append(Spacer(1, 4))
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            story.append(Paragraph(stripped[2:], h1))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], h2))
        elif stripped.startswith("### "):
            story.append(Paragraph(stripped[4:], h3))
        elif stripped.startswith("---"):
            story.append(HRFlowable(width="100%", thickness=0.5, color="#cccccc"))
            story.append(Spacer(1, 0.3*cm))
        elif stripped.startswith("> "):
            story.append(Paragraph(f"<i>{stripped[2:]}</i>", body))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            txt = stripped[2:]
            txt = _fmt_inline(txt)
            story.append(Paragraph(f"• {txt}", bullet))
        elif stripped.startswith("**") and stripped.endswith("**"):
            story.append(Paragraph(f"<b>{stripped[2:-2]}</b>", body))
        else:
            story.append(Paragraph(_fmt_inline(stripped), body))

    if in_table and table_rows:
        for row in table_rows:
            cells = [c.strip() for c in row.split("|")[1:-1]]
            story.append(Paragraph("  " + "  —  ".join(cells), body))

    return story


def _fmt_inline(text: str) -> str:
    """Convierte markdown inline a tags XML de reportlab."""
    import re
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"`([^`]+)`", r'<font face="Courier" size="9">\1</font>', text)
    return text


def email_to_pdf_bytes(content: dict, subject: str, email_date: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    except ImportError:
        raise ImportError("reportlab no instalado. Ejecutar: pip install reportlab")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
        title=subject or "Email",
    )

    styles = getSampleStyleSheet()
    header_style = ParagraphStyle("EmailHeader", parent=styles["Normal"], fontSize=9, leading=12, spaceAfter=3)
    h = content.get("headers", {})
    story = [
        Paragraph(f"<b>Subject:</b> {h.get('subject', subject) or '(sin asunto)'}", header_style),
        Paragraph(f"<b>From:</b> {h.get('from', '')}", header_style),
        Paragraph(f"<b>To:</b> {h.get('to', '')}", header_style),
        Paragraph(f"<b>Date:</b> {h.get('date', '')}", header_style),
    ]
    if h.get("cc"):
        story.append(Paragraph(f"<b>Cc:</b> {h.get('cc', '')}", header_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color="#cccccc"))
    story.append(Spacer(1, 0.5*cm))

    body_text = _plain_body(content)
    story.extend(_md_text_to_story(body_text, styles))

    doc.build(story)
    return buf.getvalue()

def email_to_pdf(path: str) -> Optional[bytes]:
    from .maildir import get_email_content
    content = get_email_content(path)
    if not content:
        return None
    h = content.get("headers", {})
    return email_to_pdf_bytes(content, h.get("subject", ""), h.get("date", ""))


# =============================================================================
# DOCX (python-docx)
# =============================================================================

def email_to_docx_bytes(content: dict, subject: str, email_date: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx no instalado. Ejecutar: pip install python-docx")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)

    h = content.get("headers", {})
    doc.add_heading(h.get("subject", subject) or "(sin asunto)", level=1)

    meta = doc.add_paragraph()
    meta.style.font.size = Pt(9)
    for label, key in [("From", "from"), ("To", "to"), ("Date", "date"), ("Cc", "cc")]:
        val = h.get(key, "")
        if val:
            meta.add_run(f"{label}: ").bold = True
            meta.add_run(f"{val}\n")

    doc.add_paragraph("─" * 60)

    for line in _plain_body(content).split("\n"):
        doc.add_paragraph(line.strip() or " ")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def email_to_docx(path: str) -> Optional[bytes]:
    from .maildir import get_email_content
    content = get_email_content(path)
    if not content:
        return None
    h = content.get("headers", {})
    return email_to_docx_bytes(content, h.get("subject", ""), h.get("date", ""))


# =============================================================================
# Dispatcher
# =============================================================================

CONVERTERS = {
    "md":   ("text/markdown",       "md",   email_to_markdown_bytes),
    "pdf":  ("application/pdf",     "pdf",  email_to_pdf_bytes),
    "docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                     "docx", email_to_docx_bytes),
}


def convert_email(filepath: str, fmt: str) -> Optional[tuple]:
    """Convierte un archivo de email al formato solicitado.

    Returns (bytes, mimetype, filename) or None.
    """
    from .maildir import get_email_content

    if fmt not in CONVERTERS:
        return None

    content = get_email_content(filepath)
    if not content:
        return None

    h = content.get("headers", {})
    subject = h.get("subject", "")
    date_str = h.get("date", "")

    mimetype, ext, converter = CONVERTERS[fmt]
    try:
        data = converter(content, subject, date_str)
    except ImportError as e:
        raise e
    except Exception as e:
        logger.error(f"Error convirtiendo a {fmt}: {e}")
        return None

    filename = _build_filename(subject, date_str, ext)
    return data, mimetype, filename
